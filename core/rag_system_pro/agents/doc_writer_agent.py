"""
doc_writer_agent.py — AI document writer that drafts polished documents
from your RAG knowledge base, templates, or free-form instructions.

Supported document types
------------------------
report       — structured analytical report with exec summary + sections
email        — professional email (subject + body)
proposal     — business / project proposal with problem, solution, timeline, budget
summary      — concise executive summary of source material
memo         — internal memo
blog_post    — blog article with intro, sections, conclusion
README       — project README with badges, install, usage, API reference stubs
cover_letter — job application cover letter
meeting_notes — formatted meeting notes with actions + owners
custom       — any free-form document type you describe

Usage
-----
from agents.doc_writer_agent import DocWriterAgent
from agents.llm_client import LLMClient

agent = DocWriterAgent(llm=LLMClient(model="deepseek-coder:6.7b"))

# Draft from a knowledge base context string (paste RAG results in)
result = agent.draft(
    doc_type="report",
    topic="Q2 sales performance",
    context="<paste retrieved chunks here>",
    instructions="Keep it under 600 words. Include a recommendations section.",
    tone="formal",
)
print(result.markdown)   # ready-to-use markdown
print(result.word_count)
result.save("Q2_report.md")   # or .docx if python-docx is installed
"""

from __future__ import annotations

import datetime
import re
import textwrap
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from .llm import LLMClient


# ── document type prompts ───────────────────────────────────────────────── #
_TYPE_PROMPTS: dict[str, str] = {
    "report": textwrap.dedent("""\
        Write a professional analytical report in Markdown.
        Structure:
        1. Executive Summary (3–5 sentences)
        2. Background / Context
        3. Findings (use sub-sections and bullet points)
        4. Analysis
        5. Recommendations
        6. Conclusion
        Use clear headings (##), tables where useful, and avoid jargon."""),

    "email": textwrap.dedent("""\
        Write a professional email. Output format:
        **Subject:** <subject line>

        <email body — greeting, body paragraphs, clear call-to-action, sign-off>
        Keep it concise and polite."""),

    "proposal": textwrap.dedent("""\
        Write a business proposal in Markdown. Include:
        1. Executive Summary
        2. Problem Statement
        3. Proposed Solution
        4. Key Benefits
        5. Implementation Timeline (table: Phase | Activities | Duration)
        6. Budget Overview (table: Item | Est. Cost)
        7. Why Us / Qualifications
        8. Next Steps
        Be persuasive but factual."""),

    "summary": textwrap.dedent("""\
        Write a concise executive summary. Rules:
        - No longer than 250 words
        - Bullet-point key takeaways at the end
        - Start with the single most important insight
        - Plain language, no fluff"""),

    "memo": textwrap.dedent("""\
        Write a formal internal memo in Markdown.
        Header block: TO / FROM / DATE / RE
        Body: Purpose, Background, Details, Action Required
        Bullet any lists. Keep it under 400 words."""),

    "blog_post": textwrap.dedent("""\
        Write a blog post in Markdown.
        Structure:
        - Compelling hook / intro paragraph
        - 3–5 body sections with ## headings
        - Practical examples or tips in each section
        - Conclusion with a call-to-action
        Tone: engaging and accessible. Aim for ~700–900 words."""),

    "README": textwrap.dedent("""\
        Write a GitHub README in Markdown. Include:
        - Project title + one-line description
        - Badges placeholder row
        - ## Features (bullet list)
        - ## Installation (code block)
        - ## Quick Start (code block)
        - ## Usage / API Reference (stub table: Function | Description | Example)
        - ## Configuration
        - ## Contributing
        - ## License
        Follow GitHub README conventions."""),

    "cover_letter": textwrap.dedent("""\
        Write a professional cover letter.
        Structure:
        - Opening: express enthusiasm, name the role
        - Body para 1: key relevant experience
        - Body para 2: specific achievements with metrics
        - Body para 3: why this company / role
        - Closing: call-to-action, thank you
        Tone: confident but not arrogant. Under 400 words."""),

    "meeting_notes": textwrap.dedent("""\
        Format meeting notes in Markdown. Include:
        ## Meeting Details
        Date / Attendees / Facilitator
        ## Agenda
        ## Discussion Summary (by agenda item)
        ## Decisions Made
        ## Action Items
        | # | Action | Owner | Due Date |
        |---|--------|-------|----------|
        ## Next Meeting"""),

    "custom": textwrap.dedent("""\
        Write the document exactly as described in the instructions and topic.
        Use Markdown formatting. Be thorough and professional."""),
}

_TONE_HINTS = {
    "formal":     "Use formal, professional language throughout.",
    "friendly":   "Use a warm, approachable tone — professional but conversational.",
    "technical":  "Use precise technical language appropriate for a specialist audience.",
    "persuasive": "Use persuasive, benefit-focused language. Lead with value.",
    "casual":     "Keep it relaxed and easy to read — like you're talking to a colleague.",
}

_SYSTEM_BASE = textwrap.dedent("""\
    You are an expert business writer and editor.
    You produce polished, well-structured documents in Markdown.
    Follow the formatting instructions exactly.
    Do NOT add meta-commentary like "Here is the document:" — start the content directly.
""")


@dataclass
class DocResult:
    doc_type: str
    topic: str
    markdown: str
    word_count: int = 0
    char_count: int = 0

    def save(self, path: str) -> str:
        """Save to .md (always) or .docx (if python-docx is available)."""
        p = Path(path)
        if p.suffix.lower() == ".docx":
            return self._save_docx(p)
        p.parent.mkdir(parents=True, exist_ok=True)
        p.with_suffix(".md").write_text(self.markdown, encoding="utf-8")
        return str(p.with_suffix(".md"))

    def _save_docx(self, p: Path) -> str:
        try:
            from docx import Document  # type: ignore
            from docx.shared import Pt  # type: ignore
        except ImportError:
            md_path = p.with_suffix(".md")
            md_path.write_text(self.markdown, encoding="utf-8")
            return f"python-docx not installed — saved as {md_path} instead"

        doc = Document()
        for line in self.markdown.splitlines():
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif re.match(r"^\d+\. ", line):
                doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
            elif line.strip():
                doc.add_paragraph(line.strip())
        p.parent.mkdir(parents=True, exist_ok=True)
        doc.save(p)
        return str(p)


class DocWriterAgent:
    """Draft any document from a topic, optional context, and instructions."""

    def __init__(self, llm: Optional[LLMClient] = None):
        self.llm = llm or LLMClient(task_type="doc")

    # ------------------------------------------------------------------ #
    def draft(
        self,
        topic: str,
        doc_type: str = "report",
        context: str = "",
        instructions: str = "",
        tone: str = "formal",
        word_limit: int = 0,
    ) -> DocResult:
        """
        Parameters
        ----------
        topic        : What the document is about.
        doc_type     : See SUPPORTED_TYPES list above.
        context      : Optional source material (RAG chunks, notes, bullet points).
        instructions : Extra formatting or content requirements.
        tone         : formal | friendly | technical | persuasive | casual
        word_limit   : 0 = no limit; otherwise appended as a constraint.
        """
        doc_type = doc_type.lower()
        type_prompt = _TYPE_PROMPTS.get(doc_type, _TYPE_PROMPTS["custom"])
        tone_hint   = _TONE_HINTS.get(tone, _TONE_HINTS["formal"])

        parts = [
            f"Document type: {doc_type}",
            f"Topic / subject: {topic}",
            f"Tone: {tone_hint}",
        ]
        if word_limit:
            parts.append(f"Word limit: approximately {word_limit} words.")
        if instructions:
            parts.append(f"Additional instructions: {instructions}")
        if context:
            parts.append(f"\n--- SOURCE MATERIAL ---\n{context}\n--- END SOURCE MATERIAL ---")

        system = _SYSTEM_BASE + "\n\n" + type_prompt
        user   = "\n".join(parts)

        markdown = self.llm.ask(system, user).strip()
        words    = len(markdown.split())
        chars    = len(markdown)
        return DocResult(
            doc_type=doc_type,
            topic=topic,
            markdown=markdown,
            word_count=words,
            char_count=chars,
        )

    # ------------------------------------------------------------------ #
    # Convenience shorthands                                               #
    # ------------------------------------------------------------------ #
    def report(self, topic: str, context: str = "", **kw) -> DocResult:
        return self.draft(topic, doc_type="report", context=context, **kw)

    def email(self, topic: str, context: str = "", **kw) -> DocResult:
        return self.draft(topic, doc_type="email", context=context, **kw)

    def proposal(self, topic: str, context: str = "", **kw) -> DocResult:
        return self.draft(topic, doc_type="proposal", context=context, **kw)

    def summary(self, source_text: str, **kw) -> DocResult:
        return self.draft("Summarise the provided material",
                           doc_type="summary", context=source_text, **kw)

    def blog_post(self, topic: str, context: str = "", **kw) -> DocResult:
        return self.draft(topic, doc_type="blog_post", context=context, **kw)

    def readme(self, project_description: str, **kw) -> DocResult:
        return self.draft(project_description, doc_type="README", **kw)

    def cover_letter(self, job_spec: str, candidate_info: str = "", **kw) -> DocResult:
        return self.draft(
            f"Cover letter for: {job_spec}",
            doc_type="cover_letter",
            context=candidate_info,
            **kw,
        )

    def meeting_notes(self, raw_notes: str, **kw) -> DocResult:
        return self.draft("Format these meeting notes",
                           doc_type="meeting_notes", context=raw_notes, **kw)


SUPPORTED_TYPES = list(_TYPE_PROMPTS.keys())
