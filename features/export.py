"""
features/export.py — Pro Export Engine  (v2.0)

Export conversations to Markdown, JSON, HTML, CSV, Excel (.xlsx),
Word (.docx), PDF, and ZIP bundles — with full file-save support,
conversation statistics, and streaming for large histories.

Pro upgrades over v1
--------------------
* ExportConfig      — single dataclass controlling titles, paths, author, etc.
* ExportStats       — computed word/token counts, role breakdown, time spans
* ExportResult      — typed return: bytes payload + suggested filename + MIME type
* ExportManager.save()  — write any format straight to disk in one call
* ExportManager.bundle() — ZIP all enabled formats into one archive
* Excel (.xlsx)     — 3-sheet workbook: Conversation (styled), Statistics, Metadata
* CSV               — flat message table with per-row word counts
* Word (.docx)      — styled headings, role paragraphs, summary appendix
* PDF               — clean layout via fpdf2 (no external LaTeX)
* Streaming         — iter_markdown() / iter_csv() yield chunks for large histories
* HTML v2           — dark/light mode toggle, copy-to-clipboard, print CSS
* Security          — HTML entity escaping, path traversal guard on save paths
* All format methods return ExportResult (breaking change from v1 str return)
  → use .text  for string formats
  → use .data  for binary formats (Excel, DOCX, PDF, ZIP)
"""

from __future__ import annotations

import csv
import html
import io
import json
import os
import re
import zipfile
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Dict, Generator, Iterator, List, Optional


# ══════════════════════════════════════════════════════════════════════════════
# Config & result types
# ══════════════════════════════════════════════════════════════════════════════

@dataclass
class ExportConfig:
    """
    Central config for all export behaviour.

    Attributes
    ----------
    title           Document / workbook title.
    author          Author name embedded in metadata.
    output_dir      Directory used by save() and bundle().
    include_stats   Append a statistics summary to supported formats.
    include_metadata Embed export metadata (date, message count, …).
    max_cell_chars  Max characters per Excel cell (Excel hard cap is 32 767).
    date_format     strftime string for timestamps in exports.
    role_labels     Custom display labels per role key.
    redact_patterns List of regex strings; matching text is replaced with ████.
    formats_in_bundle  Which formats to include in bundle().
    """
    title:               str              = "Conversation Export"
    author:              str              = "ExportManager"
    output_dir:          str              = "./exports"
    include_stats:       bool             = True
    include_metadata:    bool             = True
    max_cell_chars:      int              = 5_000
    date_format:         str              = "%Y-%m-%d %H:%M"
    role_labels:         Dict[str, str]   = field(default_factory=lambda: {
        "user":      "User",
        "assistant": "Assistant",
        "system":    "System",
    })
    redact_patterns:     List[str]        = field(default_factory=list)
    formats_in_bundle:   List[str]        = field(default_factory=lambda: [
        "markdown", "json", "html", "csv", "excel",
    ])

    def __post_init__(self) -> None:
        if self.max_cell_chars > 32_767:
            raise ValueError("max_cell_chars cannot exceed Excel's hard cap of 32 767.")


@dataclass
class ExportStats:
    """Computed statistics about the conversation."""
    total_messages:    int
    role_counts:       Dict[str, int]
    total_words:       int
    total_chars:       int
    avg_words_per_msg: float
    longest_message:   int             # chars
    shortest_message:  int             # chars
    estimated_tokens:  int             # rough: chars // 4
    export_date:       str

    def as_dict(self) -> Dict:
        return {
            "total_messages":    self.total_messages,
            "role_breakdown":    self.role_counts,
            "total_words":       self.total_words,
            "total_characters":  self.total_chars,
            "avg_words_per_msg": round(self.avg_words_per_msg, 1),
            "longest_message":   self.longest_message,
            "shortest_message":  self.shortest_message,
            "estimated_tokens":  self.estimated_tokens,
            "export_date":       self.export_date,
        }


@dataclass
class ExportResult:
    """
    Return type of all ExportManager methods.

    Attributes
    ----------
    data        Raw bytes (binary formats) or None for text formats.
    text        String content (text formats) or None for binary formats.
    filename    Suggested filename (without directory).
    mime_type   MIME type string for HTTP responses.
    stats       Computed ExportStats (if config.include_stats).
    format      Format name: "markdown" | "json" | "html" | "csv" |
                "excel" | "docx" | "pdf" | "zip"
    """
    filename:  str
    mime_type: str
    format:    str
    stats:     Optional[ExportStats] = None
    data:      Optional[bytes]       = None   # binary formats
    text:      Optional[str]         = None   # text formats

    def save_to(self, path: str) -> str:
        """Write the payload to *path* and return the resolved path."""
        path = _safe_path(path)
        os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
        if self.data is not None:
            with open(path, "wb") as f:
                f.write(self.data)
        elif self.text is not None:
            with open(path, "w", encoding="utf-8") as f:
                f.write(self.text)
        else:
            raise ValueError("ExportResult has neither data nor text to save.")
        return path

    def as_bytes(self) -> bytes:
        """
        Universal output for UI / APIs.
        Always returns bytes (Streamlit-safe).
        """
        if self.data is not None:
            return self.data
        if self.text is not None:
            return self.text.encode("utf-8")
        raise ValueError("ExportResult has no content")

    def is_binary(self) -> bool:
        return self.data is not None

    def preview(self, max_chars: int = 2000) -> str:
        """
        Safe preview for UI display.
        """
        if self.text:
            return self.text[:max_chars]
        return f"[Binary file: {self.filename}]"



# ══════════════════════════════════════════════════════════════════════════════
# Internal helpers
# ══════════════════════════════════════════════════════════════════════════════

def _safe_path(path: str) -> str:
    """Reject obvious path-traversal attempts."""
    resolved = os.path.normpath(path)
    if ".." in resolved.split(os.sep):
        raise ValueError(f"Path traversal detected in: {path!r}")
    return path


def _compute_stats(messages: List[Dict], config: ExportConfig) -> ExportStats:
    role_counts: Dict[str, int] = {}
    word_counts: List[int] = []
    char_counts: List[int] = []

    for m in messages:
        role    = m.get("role", "unknown")
        content = m.get("content", "")
        role_counts[role] = role_counts.get(role, 0) + 1
        wc = len(content.split())
        word_counts.append(wc)
        char_counts.append(len(content))

    total_words = sum(word_counts)
    total_chars = sum(char_counts)
    n           = len(messages)

    return ExportStats(
        total_messages    = n,
        role_counts       = role_counts,
        total_words       = total_words,
        total_chars       = total_chars,
        avg_words_per_msg = total_words / n if n else 0.0,
        longest_message   = max(char_counts, default=0),
        shortest_message  = min(char_counts, default=0),
        estimated_tokens  = total_chars // 4,
        export_date       = datetime.now().strftime(config.date_format),
    )


def _redact(text: str, patterns: List[str]) -> str:
    """Replace all matches of *patterns* with ████."""
    for pat in patterns:
        text = re.sub(pat, "████", text, flags=re.I)
    return text


def _slug(title: str) -> str:
    """Convert title to a filesystem-safe slug."""
    return re.sub(r"[^a-zA-Z0-9_-]", "_", title).strip("_")[:40] or "export"


def _word_count(text: str) -> int:
    return len(text.split())


# ══════════════════════════════════════════════════════════════════════════════
# ExportManager
# ══════════════════════════════════════════════════════════════════════════════

class ExportManager:
    """
    Export conversations in multiple formats.

    Quick start::

        mgr = ExportManager(ExportConfig(title="My Chat", output_dir="./out"))
        result = mgr.to_excel(messages)
        mgr.save(result)                 # writes ./out/My_Chat.xlsx

        # One-liner: save all formats as a ZIP
        mgr.save(mgr.bundle(messages))   # writes ./out/My_Chat.zip
    """

    def __init__(self, config: Optional[ExportConfig] = None) -> None:
        self._cfg = config or ExportConfig()

    # ── Convenience: save result to output_dir ─────────────────────────────

    def save(self, result: ExportResult) -> str:
        """
        Write *result* to ``config.output_dir / result.filename``.

        Returns:
            str: Full path of the written file.
        """
        out_dir = Path(self._cfg.output_dir)
        out_dir.mkdir(parents=True, exist_ok=True)
        dest = str(out_dir / result.filename)
        return result.save_to(dest)

    # ──────────────────────────────────────────────────────────────────────
    # Markdown
    # ──────────────────────────────────────────────────────────────────────

    def to_markdown(
        self,
        messages:  List[Dict],
        title:     Optional[str] = None,
    ) -> ExportResult:
        """
        Export conversation to Markdown.

        Args:
            messages: List of ``{"role": …, "content": …}`` dicts.
            title:    Override config title for this export.

        Returns:
            ExportResult with ``text`` populated.
        """
        cfg   = self._cfg
        title = title or cfg.title
        stats = _compute_stats(messages, cfg)

        parts: List[str] = [
            f"# {title}\n",
            f"_Exported on {stats.export_date}_\n",
            "---\n",
        ]

        for m in messages:
            role    = m.get("role", "unknown")
            label   = cfg.role_labels.get(role, role.capitalize())
            content = _redact(m.get("content", ""), cfg.redact_patterns)
            parts.append(f"## {label}\n\n{content}\n\n---\n")

        if cfg.include_stats:
            parts.append(self._stats_to_markdown(stats))

        text = "\n".join(parts)
        return ExportResult(
            filename  = f"{_slug(title)}.md",
            mime_type = "text/markdown; charset=utf-8",
            format    = "markdown",
            text      = text,
            stats     = stats,
        )

    @staticmethod
    def _stats_to_markdown(s: ExportStats) -> str:
        role_lines = "\n".join(f"  - {r}: {c}" for r, c in s.role_counts.items())
        return (
            "## Export Statistics\n\n"
            f"| Metric | Value |\n|--------|-------|\n"
            f"| Total Messages | {s.total_messages} |\n"
            f"| Total Words | {s.total_words:,} |\n"
            f"| Total Characters | {s.total_chars:,} |\n"
            f"| Avg Words/Message | {s.avg_words_per_msg:.1f} |\n"
            f"| Estimated Tokens | {s.estimated_tokens:,} |\n"
            f"| Longest Message | {s.longest_message:,} chars |\n\n"
            f"**Role breakdown:**\n{role_lines}\n"
        )

    # ── Streaming variant for large histories ──────────────────────────────

    def iter_markdown(self, messages: List[Dict]) -> Iterator[str]:
        """
        Yield markdown chunks one message at a time.
        Use this to stream very large conversation exports.

        Example::

            with open("out.md", "w") as f:
                for chunk in mgr.iter_markdown(messages):
                    f.write(chunk)
        """
        cfg = self._cfg
        yield f"# {cfg.title}\n\n_Exported on {datetime.now().strftime(cfg.date_format)}_\n\n---\n\n"
        for m in messages:
            role    = m.get("role", "unknown")
            label   = cfg.role_labels.get(role, role.capitalize())
            content = _redact(m.get("content", ""), cfg.redact_patterns)
            yield f"## {label}\n\n{content}\n\n---\n\n"

    # ──────────────────────────────────────────────────────────────────────
    # JSON
    # ──────────────────────────────────────────────────────────────────────

    def to_json(
        self,
        messages:  List[Dict],
        metadata:  Optional[Dict] = None,
    ) -> ExportResult:
        """
        Export conversation to JSON.

        Args:
            messages: Conversation messages.
            metadata: Extra key-value pairs embedded under ``"metadata"``.

        Returns:
            ExportResult with ``text`` populated.
        """
        cfg   = self._cfg
        stats = _compute_stats(messages, cfg)

        # Redact before serialising
        clean_messages = [
            {**m, "content": _redact(m.get("content", ""), cfg.redact_patterns)}
            for m in messages
        ]

        payload: Dict = {
            "export_date": stats.export_date,
            "title":       cfg.title,
            "author":      cfg.author,
        }
        if cfg.include_metadata:
            payload["metadata"] = metadata or {}
        if cfg.include_stats:
            payload["statistics"] = stats.as_dict()
        payload["messages"] = clean_messages

        text = json.dumps(payload, indent=2, ensure_ascii=False)
        return ExportResult(
            filename  = f"{_slug(cfg.title)}.json",
            mime_type = "application/json; charset=utf-8",
            format    = "json",
            text      = text,
            stats     = stats,
        )

    # ──────────────────────────────────────────────────────────────────────
    # HTML
    # ──────────────────────────────────────────────────────────────────────

    def to_html(
        self,
        messages:  List[Dict],
        title:     Optional[str] = None,
    ) -> ExportResult:
        """
        Export conversation to HTML with dark/light toggle, copy buttons,
        and print-optimised CSS.

        Returns:
            ExportResult with ``text`` populated.
        """
        cfg   = self._cfg
        title = title or cfg.title
        stats = _compute_stats(messages, cfg)

        bubbles: List[str] = []
        for i, m in enumerate(messages):
            role    = m.get("role", "unknown")
            label   = cfg.role_labels.get(role, role.capitalize())
            raw     = _redact(m.get("content", ""), cfg.redact_patterns)
            escaped = html.escape(raw)
            wc      = _word_count(raw)
            bubble  = (
                f'<div class="message {html.escape(role)}" id="msg-{i}">'
                f'  <div class="msg-header">'
                f'    <span class="role-label">{html.escape(label)}</span>'
                f'    <span class="msg-meta">{wc} words</span>'
                f'    <button class="copy-btn" onclick="copyMsg({i})">Copy</button>'
                f'  </div>'
                f'  <div class="msg-body" id="body-{i}">{escaped}</div>'
                f'</div>'
            )
            bubbles.append(bubble)

        stats_html = ""
        if cfg.include_stats:
            rows = "".join(
                f"<tr><td>{r}</td><td>{c}</td></tr>"
                for r, c in stats.role_counts.items()
            )
            stats_html = f"""
<section class="stats-section">
  <h2>Statistics</h2>
  <table class="stats-table">
    <tr><th>Metric</th><th>Value</th></tr>
    <tr><td>Total messages</td><td>{stats.total_messages}</td></tr>
    <tr><td>Total words</td><td>{stats.total_words:,}</td></tr>
    <tr><td>Total characters</td><td>{stats.total_chars:,}</td></tr>
    <tr><td>Avg words / message</td><td>{stats.avg_words_per_msg:.1f}</td></tr>
    <tr><td>Estimated tokens</td><td>{stats.estimated_tokens:,}</td></tr>
    {rows}
  </table>
</section>"""

        doc = f"""<!DOCTYPE html>
<html lang="en" data-theme="light">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{html.escape(title)}</title>
<style>
/* ── Reset ── */
*, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}

/* ── Tokens ── */
:root {{
  --bg:          #ffffff;
  --surface:     #f3f4f6;
  --text:        #111827;
  --text-muted:  #6b7280;
  --user-bg:     #667eea;
  --user-text:   #ffffff;
  --asst-bg:     #f3f4f6;
  --asst-text:   #111827;
  --sys-bg:      #fef9c3;
  --sys-text:    #713f12;
  --border:      #e5e7eb;
  --radius:      0.75rem;
  --shadow:      0 1px 3px rgba(0,0,0,.1);
  --font:        system-ui, -apple-system, "Segoe UI", sans-serif;
}}
[data-theme="dark"] {{
  --bg:          #111827;
  --surface:     #1f2937;
  --text:        #f9fafb;
  --text-muted:  #9ca3af;
  --user-bg:     #4f46e5;
  --user-text:   #ffffff;
  --asst-bg:     #1f2937;
  --asst-text:   #f9fafb;
  --sys-bg:      #3b3000;
  --sys-text:    #fef9c3;
  --border:      #374151;
}}

/* ── Layout ── */
body {{
  background: var(--bg);
  color: var(--text);
  font-family: var(--font);
  line-height: 1.65;
  padding: 2rem 1rem;
}}
.container {{ max-width: 820px; margin: 0 auto; }}

/* ── Header ── */
header {{ display: flex; align-items: center; justify-content: space-between;
          border-bottom: 2px solid var(--border); padding-bottom: 1rem; margin-bottom: 2rem; }}
h1 {{ font-size: 1.5rem; font-weight: 700; }}
.meta {{ color: var(--text-muted); font-size: .875rem; }}

/* ── Theme toggle ── */
.theme-btn {{
  background: var(--surface); border: 1px solid var(--border);
  color: var(--text); padding: .4rem .8rem; border-radius: 999px;
  cursor: pointer; font-size: .8rem; transition: background .2s;
}}
.theme-btn:hover {{ filter: brightness(0.92); }}

/* ── Messages ── */
.message {{
  background: var(--asst-bg); color: var(--asst-text);
  border-radius: var(--radius); margin: 1rem 0;
  box-shadow: var(--shadow); overflow: hidden;
}}
.message.user    {{ background: var(--user-bg); color: var(--user-text); }}
.message.system  {{ background: var(--sys-bg);  color: var(--sys-text);  font-style: italic; }}
.msg-header {{
  display: flex; align-items: center; gap: .5rem;
  padding: .5rem 1rem; border-bottom: 1px solid rgba(255,255,255,.15);
  font-size: .8rem;
}}
.role-label {{ font-weight: 700; flex: 1; }}
.msg-meta   {{ color: inherit; opacity: .7; }}
.copy-btn {{
  background: rgba(255,255,255,.2); border: none; color: inherit;
  padding: .2rem .6rem; border-radius: 999px; cursor: pointer; font-size: .75rem;
}}
.copy-btn:hover {{ background: rgba(255,255,255,.35); }}
.msg-body {{ padding: 1rem; white-space: pre-wrap; word-break: break-word; }}

/* ── Stats ── */
.stats-section {{ margin-top: 3rem; padding-top: 1rem; border-top: 2px solid var(--border); }}
.stats-section h2 {{ margin-bottom: 1rem; }}
.stats-table {{ width: 100%; border-collapse: collapse; font-size: .9rem; }}
.stats-table th, .stats-table td {{ padding: .5rem .75rem; border: 1px solid var(--border); text-align: left; }}
.stats-table th {{ background: var(--surface); font-weight: 600; }}

/* ── Print ── */
@media print {{
  .theme-btn, .copy-btn {{ display: none; }}
  body {{ background: white; color: black; }}
  .message {{ break-inside: avoid; }}
}}
</style>
</head>
<body>
<div class="container">
  <header>
    <div>
      <h1>{html.escape(title)}</h1>
      <div class="meta">Exported {html.escape(stats.export_date)} &nbsp;·&nbsp;
        {stats.total_messages} messages &nbsp;·&nbsp;
        {stats.total_words:,} words</div>
    </div>
    <button class="theme-btn" onclick="toggleTheme()">🌙 Dark mode</button>
  </header>

  <main>
    {"".join(bubbles)}
  </main>

  {stats_html}
</div>

<script>
function toggleTheme() {{
  const html = document.documentElement;
  const isDark = html.dataset.theme === 'dark';
  html.dataset.theme = isDark ? 'light' : 'dark';
  document.querySelector('.theme-btn').textContent = isDark ? '🌙 Dark mode' : '☀️ Light mode';
}}

function copyMsg(id) {{
  const body = document.getElementById('body-' + id);
  navigator.clipboard.writeText(body.innerText).then(() => {{
    const btn = document.querySelector(`[onclick="copyMsg(${{id}})"]`);
    const orig = btn.textContent;
    btn.textContent = '✓ Copied';
    setTimeout(() => btn.textContent = orig, 1500);
  }});
}}
</script>
</body>
</html>"""

        return ExportResult(
            filename  = f"{_slug(title)}.html",
            mime_type = "text/html; charset=utf-8",
            format    = "html",
            text      = doc,
            stats     = stats,
        )

    # ──────────────────────────────────────────────────────────────────────
    # CSV
    # ──────────────────────────────────────────────────────────────────────

    def to_csv(
        self,
        messages:  List[Dict],
        extra_columns: Optional[List[str]] = None,
    ) -> ExportResult:
        """
        Export conversation to CSV with per-message statistics.

        Columns: index, role, label, word_count, char_count, content
        Extra columns from message dicts can be requested via *extra_columns*.

        Returns:
            ExportResult with ``text`` populated.
        """
        cfg   = self._cfg
        stats = _compute_stats(messages, cfg)
        extra = extra_columns or []

        buf = io.StringIO()
        writer = csv.writer(buf, quoting=csv.QUOTE_ALL)

        header = ["index", "role", "role_label", "word_count", "char_count", "content"] + extra
        writer.writerow(header)

        for i, m in enumerate(messages):
            role    = m.get("role", "unknown")
            content = _redact(m.get("content", ""), cfg.redact_patterns)
            row = [
                i + 1,
                role,
                cfg.role_labels.get(role, role.capitalize()),
                _word_count(content),
                len(content),
                content,
            ] + [m.get(col, "") for col in extra]
            writer.writerow(row)

        return ExportResult(
            filename  = f"{_slug(cfg.title)}.csv",
            mime_type = "text/csv; charset=utf-8",
            format    = "csv",
            text      = buf.getvalue(),
            stats     = stats,
        )

    def iter_csv(self, messages: List[Dict]) -> Iterator[str]:
        """
        Yield CSV rows one line at a time for streaming large exports.

        Yields header first, then one row per message.
        """
        cfg = self._cfg
        buf = io.StringIO()
        w   = csv.writer(buf)

        w.writerow(["index", "role", "role_label", "word_count", "char_count", "content"])
        yield buf.getvalue()

        for i, m in enumerate(messages):
            buf.truncate(0); buf.seek(0)
            role    = m.get("role", "unknown")
            content = _redact(m.get("content", ""), cfg.redact_patterns)
            w.writerow([
                i + 1, role,
                cfg.role_labels.get(role, role.capitalize()),
                _word_count(content), len(content), content,
            ])
            yield buf.getvalue()

    # ──────────────────────────────────────────────────────────────────────
    # Excel (.xlsx)  — 3-sheet workbook
    # ──────────────────────────────────────────────────────────────────────

    def to_excel(
        self,
        messages:  List[Dict],
        metadata:  Optional[Dict] = None,
    ) -> ExportResult:
        """
        Export conversation to a styled 3-sheet Excel workbook.

        Sheets
        ------
        1. Conversation  — Role | Label | Words | Chars | Content
        2. Statistics    — Metrics table with role breakdown
        3. Metadata      — Config key-value pairs + optional extra metadata

        Returns:
            ExportResult with ``data`` (bytes) populated.
        """
        try:
            from openpyxl import Workbook
            from openpyxl.styles import (
                Alignment, Font, PatternFill, Border, Side,
            )
            from openpyxl.utils import get_column_letter
        except ImportError as exc:
            raise ImportError(
                "openpyxl is required for Excel export: pip install openpyxl"
            ) from exc

        cfg   = self._cfg
        stats = _compute_stats(messages, cfg)

        wb = Workbook()

        # ── Colour palette (industry-standard from SKILL.md) ──────────────
        HEADER_BG   = "1F3864"   # deep blue
        HEADER_FG   = "FFFFFF"
        USER_BG     = "DCE6F1"   # light blue
        ASST_BG     = "F2F2F2"   # light grey
        SYS_BG      = "FFFCC4"   # light yellow
        STATS_BG    = "E2EFDA"   # light green
        TITLE_FG    = "1F3864"

        thin = Side(style="thin", color="CCCCCC")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        def _header_cell(ws, row, col, value, width=20):
            c = ws.cell(row=row, column=col, value=value)
            c.font   = Font(bold=True, color=HEADER_FG, name="Arial", size=11)
            c.fill   = PatternFill("solid", start_color=HEADER_BG)
            c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            c.border = border
            ws.column_dimensions[get_column_letter(col)].width = width
            return c

        def _data_cell(ws, row, col, value, bg=None, wrap=False):
            c = ws.cell(row=row, column=col, value=value)
            c.font      = Font(name="Arial", size=10)
            c.alignment = Alignment(vertical="top", wrap_text=wrap)
            c.border    = border
            if bg:
                c.fill = PatternFill("solid", start_color=bg)
            return c

        # ── Sheet 1: Conversation ─────────────────────────────────────────
        ws1 = wb.active
        ws1.title = "Conversation"

        # Title row
        ws1.merge_cells("A1:E1")
        title_cell = ws1["A1"]
        title_cell.value     = cfg.title
        title_cell.font      = Font(bold=True, color=TITLE_FG, name="Arial", size=14)
        title_cell.alignment = Alignment(horizontal="center", vertical="center")
        ws1.row_dimensions[1].height = 28

        # Sub-header row
        ws1.merge_cells("A2:E2")
        sub = ws1["A2"]
        sub.value     = f"Exported {stats.export_date}  ·  {stats.total_messages} messages  ·  {stats.total_words:,} words"
        sub.font      = Font(italic=True, color="666666", name="Arial", size=9)
        sub.alignment = Alignment(horizontal="center")
        ws1.row_dimensions[2].height = 18

        # Column headers
        headers = [
            ("#",       5),
            ("Role",    12),
            ("Label",   12),
            ("Words",    8),
            ("Content", 80),
        ]
        for col, (h, w) in enumerate(headers, 1):
            _header_cell(ws1, 3, col, h, width=w)
        ws1.row_dimensions[3].height = 22
        ws1.freeze_panes = "A4"

        # Data rows
        _BG_MAP = {
            "user":      USER_BG,
            "assistant": ASST_BG,
            "system":    SYS_BG,
        }
        for i, m in enumerate(messages):
            r       = i + 4
            role    = m.get("role", "unknown")
            content = _redact(m.get("content", ""), cfg.redact_patterns)
            trunc   = content[: cfg.max_cell_chars]
            bg      = _BG_MAP.get(role, ASST_BG)
            wc      = _word_count(content)

            _data_cell(ws1, r, 1, i + 1, bg=bg)
            _data_cell(ws1, r, 2, role, bg=bg)
            _data_cell(ws1, r, 3, cfg.role_labels.get(role, role.capitalize()), bg=bg)
            _data_cell(ws1, r, 4, wc, bg=bg)
            _data_cell(ws1, r, 5, trunc, bg=bg, wrap=True)

            # Auto-height: approx 15pt per 80-char line, min 18
            lines = max(1, len(trunc) // 80 + trunc.count("\n"))
            ws1.row_dimensions[r].height = min(15 * lines, 400)

        # ── Sheet 2: Statistics ───────────────────────────────────────────
        ws2 = wb.create_sheet("Statistics")

        ws2.merge_cells("A1:B1")
        ws2["A1"].value = "Conversation Statistics"
        ws2["A1"].font  = Font(bold=True, color=TITLE_FG, name="Arial", size=13)
        ws2["A1"].alignment = Alignment(horizontal="center")
        ws2.row_dimensions[1].height = 24

        _header_cell(ws2, 2, 1, "Metric", width=28)
        _header_cell(ws2, 2, 2, "Value",  width=18)

        stat_rows = [
            ("Total Messages",       stats.total_messages),
            ("Total Words",          f"{stats.total_words:,}"),
            ("Total Characters",     f"{stats.total_chars:,}"),
            ("Avg Words / Message",  f"{stats.avg_words_per_msg:.1f}"),
            ("Longest Message",      f"{stats.longest_message:,} chars"),
            ("Shortest Message",     f"{stats.shortest_message:,} chars"),
            ("Estimated Tokens",     f"{stats.estimated_tokens:,}"),
            ("Export Date",          stats.export_date),
        ]
        for idx, (k, v) in enumerate(stat_rows):
            rr = idx + 3
            bg = STATS_BG if idx % 2 == 0 else None
            _data_cell(ws2, rr, 1, k, bg=bg)
            _data_cell(ws2, rr, 2, v, bg=bg)

        # Role breakdown sub-table
        offset = len(stat_rows) + 4
        _header_cell(ws2, offset, 1, "Role", width=28)
        _header_cell(ws2, offset, 2, "Message Count", width=18)
        for j, (role, count) in enumerate(stats.role_counts.items()):
            r   = offset + 1 + j
            bg  = _BG_MAP.get(role, ASST_BG)
            _data_cell(ws2, r, 1, role,  bg=bg)
            _data_cell(ws2, r, 2, count, bg=bg)

        # ── Sheet 3: Metadata ─────────────────────────────────────────────
        ws3 = wb.create_sheet("Metadata")

        ws3.merge_cells("A1:B1")
        ws3["A1"].value = "Export Metadata"
        ws3["A1"].font  = Font(bold=True, color=TITLE_FG, name="Arial", size=13)
        ws3["A1"].alignment = Alignment(horizontal="center")

        _header_cell(ws3, 2, 1, "Key",   width=28)
        _header_cell(ws3, 2, 2, "Value", width=40)

        meta_rows = [
            ("title",       cfg.title),
            ("author",      cfg.author),
            ("export_date", stats.export_date),
            ("output_dir",  cfg.output_dir),
        ]
        if metadata:
            meta_rows += list(metadata.items())

        for idx, (k, v) in enumerate(meta_rows):
            rr = idx + 3
            _data_cell(ws3, rr, 1, str(k))
            _data_cell(ws3, rr, 2, str(v))

        # ── Serialize to bytes ────────────────────────────────────────────
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)

        return ExportResult(
            filename  = f"{_slug(cfg.title)}.xlsx",
            mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            format    = "excel",
            data      = buf.read(),
            stats     = stats,
        )

    # ──────────────────────────────────────────────────────────────────────
    # Word (.docx)
    # ──────────────────────────────────────────────────────────────────────

    def to_docx(
        self,
        messages:  List[Dict],
        metadata:  Optional[Dict] = None,
    ) -> ExportResult:
        """
        Export conversation to a styled Word document.

        Requires: ``pip install python-docx``

        Returns:
            ExportResult with ``data`` (bytes) populated.
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for DOCX export: pip install python-docx"
            ) from exc

        cfg   = self._cfg
        stats = _compute_stats(messages, cfg)

        doc = Document()

        # Title
        title_p = doc.add_heading(cfg.title, level=0)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        sub_p = doc.add_paragraph(
            f"Exported {stats.export_date}  ·  {stats.total_messages} messages  ·  "
            f"{stats.total_words:,} words"
        )
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in sub_p.runs:
            run.font.size  = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run.font.italic = True

        doc.add_paragraph()

        # Role colour map
        _ROLE_COLORS: Dict[str, RGBColor] = {
            "user":      RGBColor(0x66, 0x7E, 0xEA),
            "assistant": RGBColor(0x11, 0x18, 0x27),
            "system":    RGBColor(0x71, 0x3F, 0x12),
        }

        for m in messages:
            role    = m.get("role", "unknown")
            label   = cfg.role_labels.get(role, role.capitalize())
            content = _redact(m.get("content", ""), cfg.redact_patterns)

            heading = doc.add_heading(label, level=2)
            color   = _ROLE_COLORS.get(role, RGBColor(0, 0, 0))
            for run in heading.runs:
                run.font.color.rgb = color

            body = doc.add_paragraph(content)
            body.style.font.size = Pt(11)
            doc.add_paragraph()

        # Statistics appendix
        if cfg.include_stats:
            doc.add_page_break()
            doc.add_heading("Export Statistics", level=1)
            table = doc.add_table(rows=1, cols=2)
            table.style = "Light List Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Metric"
            hdr[1].text = "Value"
            stat_rows = [
                ("Total Messages",      str(stats.total_messages)),
                ("Total Words",         f"{stats.total_words:,}"),
                ("Total Characters",    f"{stats.total_chars:,}"),
                ("Avg Words/Message",   f"{stats.avg_words_per_msg:.1f}"),
                ("Estimated Tokens",    f"{stats.estimated_tokens:,}"),
            ]
            for k, v in stat_rows:
                row = table.add_row().cells
                row[0].text = k
                row[1].text = v

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        return ExportResult(
            filename  = f"{_slug(cfg.title)}.docx",
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            format    = "docx",
            data      = buf.read(),
            stats     = stats,
        )

    # ──────────────────────────────────────────────────────────────────────
    # PDF
    # ──────────────────────────────────────────────────────────────────────

    def to_pdf(self, messages: List[Dict]) -> ExportResult:
        """
        Export conversation to PDF using fpdf2 (no LaTeX required).

        Requires: ``pip install fpdf2``

        Returns:
            ExportResult with ``data`` (bytes) populated.
        """
        try:
            from fpdf import FPDF
        except ImportError as exc:
            raise ImportError(
                "fpdf2 is required for PDF export: pip install fpdf2"
            ) from exc

        cfg   = self._cfg
        stats = _compute_stats(messages, cfg)

        _ROLE_COLORS = {
            "user":      (102, 126, 234),
            "assistant": (31,  41,  55),
            "system":    (113, 63,  18),
        }

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=18)
        pdf.add_page()

        # Title
        pdf.set_font("Helvetica", "B", 18)
        pdf.set_text_color(31, 56, 100)
        pdf.cell(0, 12, cfg.title[:80], ln=True, align="C")

        # Subtitle
        pdf.set_font("Helvetica", "I", 9)
        pdf.set_text_color(102, 102, 102)
        pdf.cell(
            0, 7,
            f"Exported {stats.export_date}  |  {stats.total_messages} messages  |  {stats.total_words:,} words",
            ln=True, align="C",
        )
        pdf.ln(6)

        for m in messages:
            role    = m.get("role", "unknown")
            label   = cfg.role_labels.get(role, role.capitalize())
            content = _redact(m.get("content", ""), cfg.redact_patterns)
            r, g, b = _ROLE_COLORS.get(role, (0, 0, 0))

            # Role heading
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(r, g, b)
            pdf.cell(0, 8, label, ln=True)

            # Message body (multi_cell handles wrapping)
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(30, 30, 30)
            # fpdf2 multi_cell requires ASCII-safe text
            safe = content.encode("latin-1", errors="replace").decode("latin-1")
            pdf.multi_cell(0, 6, safe)
            pdf.ln(4)
            pdf.set_draw_color(220, 220, 220)
            pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
            pdf.ln(4)

        # Stats page
        if cfg.include_stats:
            pdf.add_page()
            pdf.set_font("Helvetica", "B", 14)
            pdf.set_text_color(31, 56, 100)
            pdf.cell(0, 10, "Export Statistics", ln=True)
            pdf.ln(4)

            stat_rows = [
                ("Total Messages",     str(stats.total_messages)),
                ("Total Words",        f"{stats.total_words:,}"),
                ("Avg Words/Message",  f"{stats.avg_words_per_msg:.1f}"),
                ("Estimated Tokens",   f"{stats.estimated_tokens:,}"),
            ] + [(f"Role: {r}", str(c)) for r, c in stats.role_counts.items()]

            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(30, 30, 30)
            for k, v in stat_rows:
                pdf.cell(90, 7, k, border=1)
                pdf.cell(0,  7, v, border=1, ln=True)

        return ExportResult(
            filename  = f"{_slug(cfg.title)}.pdf",
            mime_type = "application/pdf",
            format    = "pdf",
            data      = bytes(pdf.output()),
            stats     = stats,
        )

    # ──────────────────────────────────────────────────────────────────────
    # ZIP bundle
    # ──────────────────────────────────────────────────────────────────────

    def bundle(
        self,
        messages:  List[Dict],
        metadata:  Optional[Dict] = None,
        formats:   Optional[List[str]] = None,
    ) -> ExportResult:
        """
        Create a ZIP archive containing multiple export formats.

        Args:
            messages: Conversation messages.
            metadata: Extra metadata passed to JSON and Excel exporters.
            formats:  Which formats to include.  Defaults to
                      ``config.formats_in_bundle``.
                      Valid values: "markdown", "json", "html", "csv",
                      "excel", "docx", "pdf".

        Returns:
            ExportResult with ``data`` (bytes of ZIP archive) populated.
        """
        cfg         = self._cfg
        requested   = formats or cfg.formats_in_bundle
        stats       = _compute_stats(messages, cfg)

        _BUILDERS = {
            "markdown": lambda: self.to_markdown(messages),
            "json":     lambda: self.to_json(messages, metadata),
            "html":     lambda: self.to_html(messages),
            "csv":      lambda: self.to_csv(messages),
            "excel":    lambda: self.to_excel(messages, metadata),
            "docx":     lambda: self.to_docx(messages, metadata),
            "pdf":      lambda: self.to_pdf(messages),
        }

        buf = io.BytesIO()
        errors: List[str] = []

        with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for fmt in requested:
                builder = _BUILDERS.get(fmt)
                if builder is None:
                    errors.append(f"Unknown format: {fmt!r}")
                    continue
                try:
                    result = builder()
                    payload = result.data if result.data is not None else (result.text or "").encode("utf-8")
                    zf.writestr(result.filename, payload)
                except Exception as exc:
                    errors.append(f"{fmt}: {exc}")

            # Manifest
            manifest = {
                "title":         cfg.title,
                "export_date":   stats.export_date,
                "formats":       requested,
                "errors":        errors,
                "statistics":    stats.as_dict(),
            }
            zf.writestr("manifest.json", json.dumps(manifest, indent=2))

        buf.seek(0)
        return ExportResult(
            filename  = f"{_slug(cfg.title)}_bundle.zip",
            mime_type = "application/zip",
            format    = "zip",
            data      = buf.read(),
            stats     = stats,
        )

    def export_all(
        self,
        messages: List[Dict],
        metadata: Optional[Dict] = None,
    ) -> Dict[str, ExportResult]:
        """
        Generate all formats at once
        """
        return {
            "markdown": self.to_markdown(messages),
            "json": self.to_json(messages, metadata),
            "html": self.to_html(messages),
            "csv": self.to_csv(messages),
            "excel": self.to_excel(messages, metadata),
            "docx": self.to_docx(messages, metadata),
            "pdf": self.to_pdf(messages),
            "zip": self.bundle(messages, metadata),
        }

    # ──────────────────────────────────────────────────────────────────────
    # Backwards-compatible static API (v1 surface)
    # ──────────────────────────────────────────────────────────────────────

    @staticmethod
    def to_markdown_str(messages: List[Dict], title: str = "Conversation") -> str:
        """v1 compat — returns a plain string (no ExportResult wrapper)."""
        return ExportManager().to_markdown(messages, title=title).text or ""

    @staticmethod
    def to_json_str(messages: List[Dict], metadata: Optional[Dict] = None) -> str:
        """v1 compat — returns a plain string (no ExportResult wrapper)."""
        return ExportManager().to_json(messages, metadata).text or ""

    @staticmethod
    def to_html_str(messages: List[Dict], title: str = "Conversation") -> str:
        """v1 compat — returns a plain string (no ExportResult wrapper)."""
        return ExportManager().to_html(messages, title=title).text or ""


__all__ = [
    "ExportConfig",
    "ExportStats",
    "ExportResult",
    "ExportManager",
]
